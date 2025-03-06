from docx import Document
from docx.shared import Inches

# Sample dictionary
data = {
    1: {'title': 'hello', 'description': 'world', 'screenshot_path': r'C:\Users\Prabal Arvind Tiwari\OneDrive\Desktop\c++\Untitled-24-02-25\screenshot_1.png'},
    2: {'title': 'hello', 'description': 'world', 'screenshot_path': r'C:\Users\Prabal Arvind Tiwari\OneDrive\Desktop\c++\Untitled-24-02-25\screenshot_2.png'},
    3: {'title': 'hello', 'description': 'world', 'screenshot_path': r'C:\Users\Prabal Arvind Tiwari\OneDrive\Desktop\c++\Untitled-24-02-25\screenshot_3.png'},
    4: {'title': 'hello', 'description': 'world', 'screenshot_path': r'C:\Users\Prabal Arvind Tiwari\OneDrive\Desktop\c++\Untitled-24-02-25\screenshot_4.png'},
    5: {'title': 'hello', 'description': 'world', 'screenshot_path': r'C:\Users\Prabal Arvind Tiwari\OneDrive\Desktop\c++\Untitled-24-02-25\screenshot_5.png'}
}

# Create Word document
doc = Document()
doc.add_heading('Report', level=1)

for key, value in data.items():
    title = value.get('title', 'No Title')
    description = value.get('description', 'No Description')
    screenshot_path = value.get('screenshot_path', '')

    # Add Title
    doc.add_heading(f"Step {key}: {title}", level=2)

    # Add Description
    doc.add_paragraph(f"Description: {description}")

    # Add Screenshot if available
    try:
        doc.add_picture(screenshot_path, width=Inches(4.0))
    except Exception as e:
        doc.add_paragraph(f"Error loading image: {e}")

    # Add Separator
    doc.add_paragraph("-" * 50)

# Save the document
file_name = 'output.docx'
doc.save(file_name)
print(f"Word document successfully created: {file_name}")
