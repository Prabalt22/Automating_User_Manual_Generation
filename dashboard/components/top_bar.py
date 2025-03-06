import flet as ft
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches

class TopBar(ft.Container):
    def __init__(self, main_instance):
        super().__init__()
        self.main_instance = main_instance
        self.date = "07-02-2025"
        self.export_button = ft.ElevatedButton(
            "Export",
            icon=ft.Icons.DOWNLOAD,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.GREEN,
                color=ft.Colors.WHITE,
            ),
            on_click=self.convertTo_pdf
        )
    
    def doc_created_alert(self, doc_extension, fileName = "ouput"):
        def cancel_delete(e):
            dialog.open = False
            self.parent.page.update()
        
        dialog = ft.AlertDialog(
            title=ft.Text("Saving......"),
            content=ft.Text(f"{doc_extension} File successfully created: {fileName} ?"),
            actions=[
                ft.TextButton("Done", on_click=cancel_delete)
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )
        
        self.main_instance.page.overlay.append(dialog)
        dialog.open = True
        self.main_instance.page.update()
    
    def convertTo_pdf(self, e):
        
        if e.control.text == "Export PDF":
            fileName = 'output.pdf'
            pdf = canvas.Canvas(fileName, pagesize=letter)
            # page dimension
            width, height = letter

            y_position = height - 50
            self.doc_created_alert("PDF", fileName)
            
            for key, value in self.main_instance.steps_data.items():
                title = value.get('title','No Title')
                description = value.get('description', 'No Description')
                screenshot_path = value.get('screenshot_path', '')
                
                # Add Title
                pdf.setFont("Helvetica-Bold", 14)
                pdf.drawString(50, y_position, f"Step {key}: {title}")
                y_position -= 20
                
                # Add description
                pdf.setFont("Helvetica", 14)
                pdf.drawString(50, y_position, f"Description: {description}")
                y_position -= 40
                
                # Add Screenshot if available
                try:
                    pdf.drawInlineImage(screenshot_path, 50, y_position - 150, width=300, height=150)
                    y_position -= 180
                except Exception as e:
                    pdf.setFillColor(colors.red)
                    pdf.drawString(50, y_position, f"Error loading image: {e}")
                    pdf.setFillColor(colors.black)
                    y_position -= 40
                    
                # add separete file
                pdf.line(30, y_position, width - 30, y_position)
                y_position -= 30
                
                # Start a new page if space is low
                if y_position < 150:
                    pdf.showPage()
                    y_position = height - 50
            pdf.save()
            print(f"PDF successfully created: {fileName}")
        elif e.control.text == "Export Word":
            # Create Word document
            doc = Document()
            doc.add_heading('Report', level=1)
            
            # Save the document
            file_name = 'output.docx'
            self.doc_created_alert("Word", file_name)

            for key, value in self.main_instance.steps_data.items():
                title = value.get('title', 'No Title')
                description = value.get('description', 'No Description')
                screenshot_path = value.get('screenshot_path', '')

                # Add Title
                doc.add_heading(f"Step {key}: {title}", level=2)

                # Add Description
                doc.add_paragraph(f"Description: {description}")

                # Add Screenshot if available
                try:
                    doc.add_picture(screenshot_path, width=Inches(4.0))
                except Exception as e:
                    doc.add_paragraph(f"Error loading image: {e}")

                # Add Separator
                doc.add_paragraph("-" * 50)

            doc.save(file_name)
            print(f"Word document successfully created: {file_name}")
        else:
            print(e.control)
        
        
    def set_exportBtn_label(self, e):
        self.export_button.text = f"Export {e.control.value}"
        self.export_button.update()
    
    def build(self):  
        self.content = ft.Row(
            [
                ft.IconButton(icon=ft.Icons.MENU, tooltip="Menu"),
                ft.Text("Untitled - " + self.date, size=16, weight=ft.FontWeight.BOLD),
                ft.Row(
                    [
                        ft.Dropdown(
                            width=100,
                            options=[
                                ft.dropdown.Option("PDF"),
                                ft.dropdown.Option("Word"),
                            ],
                            on_change=self.set_exportBtn_label,
                        ),
                        self.export_button
                    ],
                    spacing=10,
                )
            ],
            alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
        )
        self.padding=10
        self.bgcolor=ft.Colors.WHITE
        self.border=ft.border.only(bottom=ft.BorderSide(1, ft.Colors.BLACK12))
        